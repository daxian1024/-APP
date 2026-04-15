from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches

ROOT = Path(r"c:\Users\31849\Desktop\互联网+智慧护理移动护理平台APP")
IMG_DIR = ROOT / "docs" / "images"
OUT = ROOT / "5.系统架构与详细设计说明书.docx"

IMG_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size=24):
    cands = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for f in cands:
        if Path(f).exists():
            try:
                return ImageFont.truetype(f, size)
            except Exception:
                pass
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill, outline, font):
    draw.rounded_rectangle(xy, radius=12, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    w, h = draw.textbbox((0, 0), text, font=font)[2:]
    tx = x1 + (x2 - x1 - w) / 2
    ty = y1 + (y2 - y1 - h) / 2
    draw.text((tx, ty), text, fill=(30, 50, 90), font=font)


def make_img_2_1(path):
    img = Image.new("RGB", (1400, 900), (248, 251, 255))
    d = ImageDraw.Draw(img)
    f_title = get_font(42)
    f = get_font(24)
    d.text((460, 20), "图2-1 系统用例视图", fill=(25, 55, 105), font=f_title)
    d.rectangle((80, 90, 1320, 840), outline=(70, 110, 180), width=3)

    draw_box(d, (280, 160, 520, 230), "注册/登录", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (580, 160, 860, 230), "浏览服务项目", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (920, 160, 1200, 230), "下单与支付", (255, 243, 210), (200, 168, 90), f)

    draw_box(d, (280, 300, 520, 370), "地址管理", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (580, 300, 860, 370), "订单查询", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (920, 300, 1200, 370), "评价与投诉", (255, 243, 210), (200, 168, 90), f)

    draw_box(d, (480, 500, 780, 570), "执行服务并更新状态", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (840, 500, 1200, 570), "查看分配订单", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (480, 660, 760, 730), "分配护士", (255, 243, 210), (200, 168, 90), f)
    draw_box(d, (820, 660, 1240, 730), "服务项目CRUD与通知", (255, 243, 210), (200, 168, 90), f)

    d.text((120, 180), "老人用户", fill=(40, 60, 95), font=get_font(30))
    d.text((120, 520), "护士", fill=(40, 60, 95), font=get_font(30))
    d.text((120, 680), "管理员", fill=(40, 60, 95), font=get_font(30))

    img.save(path)


def make_img_2_2(path):
    img = Image.new("RGB", (1400, 900), (250, 252, 255))
    d = ImageDraw.Draw(img)
    d.text((430, 20), "图2-2 层次结构（项目结构图）", fill=(25, 55, 105), font=get_font(42))
    f = get_font(26)

    draw_box(d, (120, 110, 1280, 220), "表示层：frontend/index.html · frontend/admin.html", (232, 240, 255), (95, 130, 200), f)
    draw_box(d, (120, 260, 1280, 390), "应用层：auth/services/orders/addresses/admin/notifications", (235, 247, 239), (95, 160, 120), f)
    draw_box(d, (120, 430, 1280, 540), "服务层：RBAC、签名校验、限流防刷、统一响应", (255, 244, 225), (205, 150, 90), f)
    draw_box(d, (120, 580, 1280, 730), "数据层：MySQL、Redis、Kafka/Flink、Hive/ClickHouse", (242, 236, 255), (130, 110, 190), f)

    d.line((700, 220, 700, 260), fill=(90, 120, 180), width=4)
    d.line((700, 390, 700, 430), fill=(90, 120, 180), width=4)
    d.line((700, 540, 700, 580), fill=(90, 120, 180), width=4)

    img.save(path)


def make_img_2_3(path):
    img = Image.new("RGB", (1400, 900), (247, 250, 255))
    d = ImageDraw.Draw(img)
    d.text((520, 20), "图2-3 部署视图", fill=(25, 55, 105), font=get_font(42))
    f = get_font(24)

    draw_box(d, (100, 140, 360, 250), "移动端/H5", (232, 240, 255), (95, 130, 200), f)
    draw_box(d, (420, 140, 680, 250), "Nginx 网关", (232, 240, 255), (95, 130, 200), f)
    draw_box(d, (740, 140, 1020, 250), "Flask API", (232, 240, 255), (95, 130, 200), f)

    draw_box(d, (120, 360, 360, 470), "MySQL", (235, 247, 239), (95, 160, 120), f)
    draw_box(d, (420, 360, 680, 470), "Redis", (235, 247, 239), (95, 160, 120), f)
    draw_box(d, (760, 360, 1080, 470), "Kafka + Flink", (255, 244, 225), (205, 150, 90), f)

    draw_box(d, (280, 580, 620, 700), "Hive/HDFS", (242, 236, 255), (130, 110, 190), f)
    draw_box(d, (700, 580, 1080, 700), "ClickHouse/BI", (242, 236, 255), (130, 110, 190), f)

    d.line((360, 195, 420, 195), fill=(90, 120, 180), width=4)
    d.line((680, 195, 740, 195), fill=(90, 120, 180), width=4)
    d.line((880, 250, 880, 360), fill=(90, 120, 180), width=4)
    d.line((850, 470, 850, 580), fill=(90, 120, 180), width=4)

    img.save(path)


def make_img_3_1(path):
    img = Image.new("RGB", (1400, 900), (248, 251, 255))
    d = ImageDraw.Draw(img)
    d.text((380, 20), "图3-1 系统功能结构图（智慧护理平台）", fill=(25, 55, 105), font=get_font(42))
    f = get_font(24)

    draw_box(d, (560, 100, 860, 180), "智慧护理平台", (65, 110, 210), (65, 110, 210), f)

    draw_box(d, (80, 260, 320, 340), "用户与认证", (223, 234, 255), (110, 145, 220), f)
    draw_box(d, (360, 260, 600, 340), "服务项目", (223, 234, 255), (110, 145, 220), f)
    draw_box(d, (640, 260, 900, 340), "订单与支付", (223, 234, 255), (110, 145, 220), f)
    draw_box(d, (940, 260, 1220, 340), "评价与投诉", (223, 234, 255), (110, 145, 220), f)

    draw_box(d, (140, 450, 430, 530), "管理端（护士/管理员）", (232, 246, 237), (100, 160, 120), f)
    draw_box(d, (490, 450, 830, 530), "安全审计（RBAC/签名/限流）", (232, 246, 237), (100, 160, 120), f)
    draw_box(d, (890, 450, 1220, 530), "通知中心", (232, 246, 237), (100, 160, 120), f)

    draw_box(d, (140, 640, 430, 740), "离线分析\npandas/numpy/PySpark", (255, 244, 225), (205, 150, 90), f)
    draw_box(d, (490, 640, 830, 740), "实时计算\nKafka/Flink", (255, 244, 225), (205, 150, 90), f)
    draw_box(d, (890, 640, 1220, 740), "AI服务\n推荐/预警/分诊", (255, 244, 225), (205, 150, 90), f)

    img.save(path)


def build_docx():
    p1 = IMG_DIR / "图2-1_系统用例视图.png"
    p2 = IMG_DIR / "图2-2_层次结构图.png"
    p3 = IMG_DIR / "图2-3_部署视图.png"
    p4 = IMG_DIR / "图3-1_系统功能结构图.png"

    make_img_2_1(p1)
    make_img_2_2(p2)
    make_img_2_3(p3)
    make_img_3_1(p4)

    doc = Document()
    doc.add_heading("5. 系统架构与详细设计说明书", level=1)

    doc.add_heading("1. 简介", level=2)
    doc.add_paragraph("1.1 目的：说明平台架构、模块设计、接口规范与测试方案，为开发与运维提供依据。")
    doc.add_paragraph("1.2 范围：覆盖老人端、护士/管理员端、Flask后端、数据分析与AI扩展。")
    doc.add_paragraph("1.3 术语：RBAC、JWT、CDC、ODS/DWD/DWS、DLP。")

    doc.add_heading("2. 架构设计", level=2)
    doc.add_paragraph("2.1 数据处理架构：MySQL业务库 + Redis缓存，实时链路MySQL→Kafka→Flink→Hive/ClickHouse。")
    doc.add_paragraph("2.2 数据分析架构：离线使用pandas/numpy/PySpark，实时使用Flink窗口告警。")
    doc.add_paragraph("2.3 数据安全架构：JWT认证、RBAC权限、签名校验、限流防刷、审计日志。")
    doc.add_paragraph("2.4 Web架构：前后端分离，Nginx网关，Flask API服务。")

    doc.add_picture(str(p1), width=Inches(6.2))
    doc.add_paragraph("图2-1 系统用例视图")
    doc.add_picture(str(p2), width=Inches(6.2))
    doc.add_paragraph("图2-2 层次结构图")
    doc.add_picture(str(p3), width=Inches(6.2))
    doc.add_paragraph("图2-3 部署视图")

    doc.add_heading("3. 详细设计", level=2)
    doc.add_paragraph("3.1 模块：认证、服务项目、订单流转、支付、通知、安全、数据与AI。")

    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "模块"
    hdr[1].text = "编号"
    hdr[2].text = "说明"
    hdr[3].text = "关键接口"

    rows = [
        ("认证模块", "AUTH-01", "注册/登录/当前用户", "/api/auth/register /login /me"),
        ("服务模块", "SVC-02", "服务查询/热门缓存/CRUD", "/api/services /api/services/popular"),
        ("订单模块", "ORD-03", "下单/查询/状态流转/时间线", "/api/orders /api/admin/orders"),
        ("支付模块", "PAY-04", "模拟支付", "/api/orders/{id}/pay"),
        ("通知模块", "NTF-05", "站内/短信/微信渠道通知", "/api/admin/notifications/send"),
        ("安全模块", "SEC-06", "RBAC/签名/限流/审计", "permission_required verify_signature rate_limit"),
    ]
    for r in rows:
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text, c[3].text = r

    doc.add_paragraph("3.3 接口设计示例：")
    doc.add_paragraph("- 登录：POST /api/auth/login")
    doc.add_paragraph("- 状态流转：PATCH /api/admin/orders/{order_id}/status（需签名）")
    doc.add_paragraph("- 服务新增：POST /api/services（管理员）")

    doc.add_picture(str(p4), width=Inches(6.2))
    doc.add_paragraph("图3-1 系统功能结构图")

    doc.add_heading("4. 附录", level=2)
    doc.add_paragraph("附图已以内嵌形式写入本DOCX，同时在 docs/images 目录保留PNG源文件。")

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build_docx()
